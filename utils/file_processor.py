import os
import json
import tempfile
import mimetypes
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union
import logging
from datetime import datetime
import shutil
import hashlib
import base64

try:
    from PIL import Image, ImageOps
    import pytesseract
    HAS_IMAGE_DEPS = True
except ImportError:
    HAS_IMAGE_DEPS = False
    Image = None
    pytesseract = None

try:
    import PyPDF2
    HAS_PDF_DEPS = True
except ImportError:
    HAS_PDF_DEPS = False
    PyPDF2 = None

try:
    import docx
    HAS_DOCX_DEPS = True
except ImportError:
    HAS_DOCX_DEPS = False
    docx = None

logger = logging.getLogger(__name__)

class FileProcessor:
    """文件处理器"""
    
    def __init__(self, upload_dir: str = "uploads", max_file_size: int = 10 * 1024 * 1024):
        """
        初始化文件处理器
        
        Args:
            upload_dir: 上传文件目录
            max_file_size: 最大文件大小（字节）
        """
        self.upload_dir = Path(upload_dir)
        self.max_file_size = max_file_size
        self.create_upload_dir()
        
        # 支持的文件类型
        self.supported_types = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            'document': ['.pdf', '.doc', '.docx', '.txt', '.md', '.rtf'],
            'audio': ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.aac'],
            'video': ['.mp4', '.avi', '.mov', '.mkv', '.webm']
        }
    
    def create_upload_dir(self):
        """创建上传目录"""
        try:
            self.upload_dir.mkdir(exist_ok=True, parents=True)
            
            # 创建子目录
            for subdir in ['images', 'documents', 'audio', 'video', 'temp']:
                (self.upload_dir / subdir).mkdir(exist_ok=True)
            
            logger.info(f"上传目录已创建: {self.upload_dir}")
        except Exception as e:
            logger.error(f"创建上传目录失败: {e}")
            raise
    
    def get_file_info(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        获取文件信息
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件信息字典
        """
        try:
            path = Path(file_path)
            
            if not path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 获取文件统计信息
            stat = path.stat()
            
            # 确定文件类型
            mime_type, _ = mimetypes.guess_type(str(path))
            file_type = self._categorize_file(path)
            
            # 计算文件哈希
            file_hash = self._calculate_file_hash(path)
            
            return {
                "name": path.name,
                "path": str(path.absolute()),
                "size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "type": file_type,
                "mime_type": mime_type or "application/octet-stream",
                "extension": path.suffix.lower(),
                "hash": file_hash
            }
            
        except Exception as e:
            logger.error(f"获取文件信息失败: {e}")
            return {}
    
    def _categorize_file(self, file_path: Path) -> str:
        """
        分类文件类型
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件类型: 'image', 'document', 'audio', 'video', 'other'
        """
        ext = file_path.suffix.lower()
        
        for file_type, extensions in self.supported_types.items():
            if ext in extensions:
                return file_type
        
        return 'other'
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """
        计算文件哈希值
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件哈希值
        """
        try:
            hasher = hashlib.sha256()
            with open(file_path, 'rb') as f:
                while chunk := f.read(8192):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"计算文件哈希失败: {e}")
            return ""
    
    def save_uploaded_file(self, file_content: bytes, filename: str, 
                          category: str = "temp") -> Optional[Path]:
        """
        保存上传的文件
        
        Args:
            file_content: 文件内容字节
            filename: 原始文件名
            category: 文件分类
        
        Returns:
            保存的文件路径，失败则返回None
        """
        try:
            # 检查文件大小
            if len(file_content) > self.max_file_size:
                raise ValueError(f"文件大小超过限制: {len(file_content)} > {self.max_file_size}")
            
            # 生成安全的文件名
            safe_filename = self._make_safe_filename(filename)
            
            # 确定保存目录
            if category in self.supported_types:
                save_dir = self.upload_dir / category
            else:
                save_dir = self.upload_dir / "temp"
            
            # 确保目录存在
            save_dir.mkdir(exist_ok=True)
            
            # 生成唯一文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_name = f"{timestamp}_{safe_filename}"
            save_path = save_dir / unique_name
            
            # 保存文件
            with open(save_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"文件保存成功: {save_path}")
            return save_path
            
        except Exception as e:
            logger.error(f"保存文件失败: {e}")
            return None
    
    def _make_safe_filename(self, filename: str) -> str:
        """
        生成安全的文件名
        
        Args:
            filename: 原始文件名
        
        Returns:
            安全的文件名
        """
        # 移除路径信息
        safe_name = Path(filename).name
        
        # 移除不安全的字符
        safe_name = "".join(c for c in safe_name if c.isalnum() or c in "._- ")
        
        # 限制长度
        if len(safe_name) > 100:
            name, ext = os.path.splitext(safe_name)
            safe_name = name[:95] + ext
        
        return safe_name
    
    def image_to_text(self, image_path: Union[str, Path]) -> str:
        """
        图片转文本（OCR）
        
        Args:
            image_path: 图片文件路径
        
        Returns:
            识别的文本
        """
        if not HAS_IMAGE_DEPS:
            logger.error("未安装图片处理依赖，请运行: pip install Pillow pytesseract")
            return ""
        
        try:
            image = Image.open(image_path)
            
            # 预处理图片
            image = self._preprocess_image(image)
            
            # OCR识别
            text = pytesseract.image_to_string(
                image,
                lang='chi_sim+eng'  # 中文简体+英文
            )
            
            logger.info(f"图片OCR完成: {len(text)} 字符")
            return text.strip()
            
        except Exception as e:
            logger.error(f"图片OCR失败: {e}")
            return ""
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """
        预处理图片以改善OCR效果
        
        Args:
            image: PIL图片对象
        
        Returns:
            预处理后的图片
        """
        # 转换为灰度图
        if image.mode != 'L':
            image = image.convert('L')
        
        # 增强对比度
        image = ImageOps.autocontrast(image, cutoff=2)
        
        # 调整大小（如果太大）
        max_size = 2000
        if max(image.size) > max_size:
            ratio = max_size / max(image.size)
            new_size = tuple(int(dim * ratio) for dim in image.size)
            image = image.resize(new_size, Image.Resampling.LANCZOS)
        
        return image
    
    def pdf_to_text(self, pdf_path: Union[str, Path]) -> str:
        """
        PDF转文本
        
        Args:
            pdf_path: PDF文件路径
        
        Returns:
            提取的文本
        """
        if not HAS_PDF_DEPS:
            logger.error("未安装PDF处理依赖，请运行: pip install PyPDF2")
            return ""
        
        try:
            text_parts = []
            
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            result = "\n".join(text_parts)
            logger.info(f"PDF提取完成: {len(result)} 字符，{len(pdf_reader.pages)} 页")
            return result
            
        except Exception as e:
            logger.error(f"PDF提取失败: {e}")
            return ""
    
    def docx_to_text(self, docx_path: Union[str, Path]) -> str:
        """
        DOCX转文本
        
        Args:
            docx_path: DOCX文件路径
        
        Returns:
            提取的文本
        """
        if not HAS_DOCX_DEPS:
            logger.error("未安装DOCX处理依赖，请运行: pip install python-docx")
            return ""
        
        try:
            doc = docx.Document(docx_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            result = "\n".join(text_parts)
            logger.info(f"DOCX提取完成: {len(result)} 字符")
            return result
            
        except Exception as e:
            logger.error(f"DOCX提取失败: {e}")
            return ""
    
    def extract_text_from_file(self, file_path: Union[str, Path]) -> str:
        """
        从文件中提取文本
        
        Args:
            file_path: 文件路径
        
        Returns:
            提取的文本
        """
        try:
            path = Path(file_path)
            ext = path.suffix.lower()
            
            if not path.exists():
                return f"文件不存在: {file_path}"
            
            if ext in ['.txt', '.md', '.rtf']:
                # 文本文件直接读取
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
                # 图片文件OCR
                return self.image_to_text(path)
            
            elif ext == '.pdf':
                # PDF文件
                return self.pdf_to_text(path)
            
            elif ext in ['.doc', '.docx']:
                # Word文档
                return self.docx_to_text(path)
            
            else:
                # 其他文件类型尝试读取
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
                
        except Exception as e:
            logger.error(f"提取文件文本失败: {e}")
            return f"无法提取文件内容: {str(e)}"
    
    def compress_image(self, image_path: Union[str, Path], 
                      max_size: Tuple[int, int] = (1920, 1080),
                      quality: int = 85) -> Optional[Path]:
        """
        压缩图片
        
        Args:
            image_path: 图片路径
            max_size: 最大尺寸 (宽, 高)
            quality: JPEG质量 (1-100)
        
        Returns:
            压缩后的图片路径
        """
        if not HAS_IMAGE_DEPS:
            logger.error("未安装图片处理依赖")
            return None
        
        try:
            image = Image.open(image_path)
            
            # 调整大小
            image.thumbnail(max_size, Image.Resampling.LANCZOS)
            
            # 保存压缩后的图片
            compressed_path = Path(image_path).with_stem(
                Path(image_path).stem + "_compressed"
            )
            
            # 根据格式保存
            if image_path.lower().endswith(('.jpg', '.jpeg')):
                image.save(compressed_path, 'JPEG', quality=quality, optimize=True)
            elif image_path.lower().endswith('.png'):
                image.save(compressed_path, 'PNG', optimize=True)
            else:
                # 转换为JPEG保存
                if image.mode in ('RGBA', 'LA', 'P'):
                    # 转换RGBA为RGB
                    rgb_image = Image.new('RGB', image.size, (255, 255, 255))
                    rgb_image.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                    image = rgb_image
                
                compressed_path = compressed_path.with_suffix('.jpg')
                image.save(compressed_path, 'JPEG', quality=quality, optimize=True)
            
            original_size = Path(image_path).stat().st_size
            compressed_size = compressed_path.stat().st_size
            reduction = (1 - compressed_size / original_size) * 100
            
            logger.info(f"图片压缩完成: {original_size:,} -> {compressed_size:,} "
                       f"字节 ({reduction:.1f}% 减少)")
            
            return compressed_path
            
        except Exception as e:
            logger.error(f"图片压缩失败: {e}")
            return None
    
    def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        提取文件元数据
        
        Args:
            file_path: 文件路径
        
        Returns:
            元数据字典
        """
        try:
            path = Path(file_path)
            file_info = self.get_file_info(path)
            metadata = {"basic_info": file_info}
            
            # 根据文件类型提取额外元数据
            if file_info["type"] == "image" and HAS_IMAGE_DEPS:
                metadata["image_info"] = self._extract_image_metadata(path)
            
            elif file_info["type"] == "document":
                if path.suffix.lower() == '.pdf' and HAS_PDF_DEPS:
                    metadata["pdf_info"] = self._extract_pdf_metadata(path)
                elif path.suffix.lower() in ['.doc', '.docx'] and HAS_DOCX_DEPS:
                    metadata["doc_info"] = self._extract_docx_metadata(path)
            
            return metadata
            
        except Exception as e:
            logger.error(f"提取元数据失败: {e}")
            return {}
    
    def _extract_image_metadata(self, image_path: Path) -> Dict[str, Any]:
        """提取图片元数据"""
        try:
            image = Image.open(image_path)
            
            return {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
                "info": image.info
            }
        except:
            return {}
    
    def _extract_pdf_metadata(self, pdf_path: Path) -> Dict[str, Any]:
        """提取PDF元数据"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                return {
                    "pages": len(pdf_reader.pages),
                    "metadata": pdf_reader.metadata
                }
        except:
            return {}
    
    def _extract_docx_metadata(self, docx_path: Path) -> Dict[str, Any]:
        """提取DOCX元数据"""
        try:
            doc = docx.Document(docx_path)
            
            return {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "core_properties": {
                    "title": doc.core_properties.title,
                    "author": doc.core_properties.author,
                    "created": doc.core_properties.created,
                    "modified": doc.core_properties.modified
                }
            }
        except:
            return {}
    
    def cleanup_temp_files(self, max_age_hours: int = 24):
        """
        清理临时文件
        
        Args:
            max_age_hours: 最大保存时间（小时）
        """
        try:
            temp_dir = self.upload_dir / "temp"
            if not temp_dir.exists():
                return
            
            cutoff_time = datetime.now().timestamp() - (max_age_hours * 3600)
            
            for file_path in temp_dir.iterdir():
                if file_path.is_file():
                    file_age = file_path.stat().st_mtime
                    if file_age < cutoff_time:
                        try:
                            file_path.unlink()
                            logger.debug(f"已删除临时文件: {file_path}")
                        except Exception as e:
                            logger.error(f"删除临时文件失败 {file_path}: {e}")
            
            logger.info(f"临时文件清理完成: {temp_dir}")
            
        except Exception as e:
            logger.error(f"清理临时文件失败: {e}")

# 全局文件处理器实例
file_processor = FileProcessor()

# 兼容性函数
def image_to_text(image_path: str) -> str:
    """图片转文本（兼容性函数）"""
    return file_processor.image_to_text(image_path)

def extract_text_from_file(file_path: str) -> str:
    """从文件提取文本（兼容性函数）"""
    return file_processor.extract_text_from_file(file_path)